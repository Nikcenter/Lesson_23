from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Pt


# Якщо ви на Windows, потрібно вказати шлях до tesseract.exe
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def pdf_to_word_ocr(pdf_file, output_docx='output_document.docx', lang='ukr'):
    """
    Конвертує сканований PDF у Word-документ за допомогою OCR.
    pdf_file: шлях до PDF-файлу
    output_docx: вихідний Word-файл
    lang: мова OCR ('ukr', 'eng', 'rus' тощо)
    """
    # Крок 1: Створюємо новий Word-документ
    doc = Document()

    # Налаштуємо базовий стиль (не обов’язково)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Крок 2: Конвертуємо кожну сторінку PDF у зображення з якістю 300 dpi
    pages = convert_from_path(pdf_file, 300)

    # Крок 3: Виконуємо OCR для кожної сторінки і зберігаємо текст у документ
    for page_number, page_image in enumerate(pages):
        # Отримуємо текст зображення
        text = pytesseract.image_to_string(page_image, lang=lang)

        # Додаємо нові рядки після крапки, якщо перед нею літера
        cleaned_text = ''
        for line in text.splitlines():
            if line.strip():
                processed_line = ''.join(
                    [f"{char}\n" if char == '.' and i > 0 and line[i - 1].isalpha() else char for i, char in
                     enumerate(line)])
                cleaned_text += processed_line + ' '

        # Друкуємо (для відладки)
        print(f"Текст зі сторінки {page_number + 1}:")
        print(cleaned_text)

        # Додаємо текст у Word як єдиний параграф
        doc.add_paragraph(f"--- Сторінка {page_number + 1} ---")
        doc.add_paragraph(cleaned_text)

    # Крок 4: Зберігаємо документ
    doc.save(output_docx)
    print(f"Документ збережено як: {output_docx}")


# Виклик функції для виконання перетворення
pdf_to_word_ocr('document.pdf', 'scanned_to_word.docx', 'ukr')
